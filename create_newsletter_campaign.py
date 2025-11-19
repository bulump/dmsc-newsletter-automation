#!/usr/bin/env python3
"""
DMSC Newsletter Campaign Automation
Phase 3: Complete automation from Dropbox → WIX → MailChimp

This script:
1. Finds newsletter PDF in Dropbox by month name
2. Creates Dropbox share link
3. Extracts meeting info from Ted's Thoughts document
4. Imports PDF to WIX Media Manager
5. Creates CMS entry in WIX Newsletters collection
6. Creates MailChimp campaign with WIX URL
7. Leaves campaign as DRAFT for manual review

Phase 2 (future): Will auto-send after review
"""

import requests
import json
import sys
import os
import io
from datetime import datetime
from pathlib import Path
from docx import Document

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# MailChimp API Configuration
MAILCHIMP_API_KEY = os.getenv("MAILCHIMP_API_KEY")
MAILCHIMP_LIST_ID = os.getenv("MAILCHIMP_LIST_ID")

# WIX API Configuration
WIX_API_KEY = os.getenv("WIX_API_KEY")
WIX_SITE_ID = os.getenv("WIX_SITE_ID")

# Dropbox API Configuration
DROPBOX_ACCESS_TOKEN = os.getenv("DROPBOX_ACCESS_TOKEN")

# Validate all required credentials
required_vars = {
    "MAILCHIMP_API_KEY": MAILCHIMP_API_KEY,
    "MAILCHIMP_LIST_ID": MAILCHIMP_LIST_ID,
    "WIX_API_KEY": WIX_API_KEY,
    "WIX_SITE_ID": WIX_SITE_ID,
    "DROPBOX_ACCESS_TOKEN": DROPBOX_ACCESS_TOKEN
}

for var_name, var_value in required_vars.items():
    if not var_value:
        print(f"Error: {var_name} not found in environment variables")
        print("Please check your .env file")
        sys.exit(1)

# Extract data center from API key
DC = MAILCHIMP_API_KEY.split("-")[-1]
MAILCHIMP_BASE_URL = f"https://{DC}.api.mailchimp.com/3.0"

# Template file location (in same directory as script)
SCRIPT_DIR = Path(__file__).parent
TEMPLATE_FILE = SCRIPT_DIR / "newsletter_template.html"

# Authentication
MAILCHIMP_AUTH = ("anystring", MAILCHIMP_API_KEY)
HEADERS = {"Content-Type": "application/json"}


# ============================================================
# DROPBOX FUNCTIONS
# ============================================================

def find_newsletter_pdf(month, year=None):
    """Find the newsletter PDF in Dropbox for a given month"""
    if not year:
        year = datetime.now().year

    folder_path = f"/Newsletter/Monthly Newsletters/{year} Newsletter/{month}"
    print(f"  Looking in: {folder_path}")

    url = "https://api.dropboxapi.com/2/files/list_folder"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {"path": folder_path}

    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        entries = response.json().get("entries", [])
        pdf_files = [e for e in entries if e["name"].endswith("_Web.pdf")]

        if pdf_files:
            # Return both path and actual filename with correct casing
            return {
                "path": pdf_files[0]["path_lower"],
                "filename": pdf_files[0]["name"]
            }

    return None


def create_dropbox_share_link(file_path):
    """Create a share link for a Dropbox file"""
    url = "https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {
        "path": file_path,
        "settings": {"requested_visibility": "public"}
    }

    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        share_url = response.json()["url"]
        # Convert to direct download URL
        return share_url.replace("dl=0", "dl=1").replace("www.dropbox.com", "dl.dropboxusercontent.com")
    elif response.status_code == 409:
        # Link already exists, get it
        get_url = "https://api.dropboxapi.com/2/sharing/list_shared_links"
        get_data = {"path": file_path}
        get_response = requests.post(get_url, headers=headers, json=get_data)

        if get_response.status_code == 200:
            links = get_response.json().get("links", [])
            if links:
                share_url = links[0]["url"]
                return share_url.replace("dl=0", "dl=1").replace("www.dropbox.com", "dl.dropboxusercontent.com")

    return None


def find_teds_thoughts(month, year=None):
    """Find Ted's Thoughts document in Dropbox"""
    if not year:
        year = datetime.now().year

    folder_path = f"/Newsletter/Monthly Newsletters/{year} Newsletter/{month}"

    url = "https://api.dropboxapi.com/2/files/list_folder"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {"path": folder_path}

    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        entries = response.json().get("entries", [])
        ted_files = [e for e in entries if "ted" in e["name"].lower() and e["name"].endswith(".docx")]

        if ted_files:
            return ted_files[0]["path_lower"]

    return None


def download_dropbox_file(file_path):
    """Download a file from Dropbox"""
    url = "https://content.dropboxapi.com/2/files/download"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Dropbox-API-Arg": f'{{"path": "{file_path}"}}'
    }

    response = requests.post(url, headers=headers)

    if response.status_code == 200:
        return response.content

    return None


def extract_meeting_info(docx_bytes):
    """Extract meeting information from Ted's Thoughts document"""
    try:
        doc = Document(io.BytesIO(docx_bytes))

        # Read first few paragraphs looking for meeting information
        meeting_info = []
        for para in doc.paragraphs[:10]:
            text = para.text.strip()
            if text and ("meeting" in text.lower() or "speaker" in text.lower() or "program" in text.lower()):
                meeting_info.append(text)

        if meeting_info:
            return meeting_info[0]
        else:
            # Fallback: return first non-empty paragraph
            for para in doc.paragraphs[:5]:
                if para.text.strip():
                    return para.text.strip()

        return "See newsletter for meeting details"

    except Exception as e:
        print(f"  Warning: Could not extract meeting info - {e}")
        return "See newsletter for meeting details"


# ============================================================
# WIX FUNCTIONS
# ============================================================

def import_file_to_wix(file_url, display_name):
    """Import file to WIX Media Manager from URL"""
    endpoint = "https://www.wixapis.com/site-media/v1/files/import"

    headers = {
        "Authorization": WIX_API_KEY,
        "wix-site-id": WIX_SITE_ID,
        "Content-Type": "application/json"
    }

    data = {
        "url": file_url,
        "mimeType": "application/pdf",
        "displayName": display_name
    }

    response = requests.post(endpoint, headers=headers, json=data)

    if response.status_code == 200:
        result = response.json()
        file_info = result.get("file", {})
        file_id = file_info.get("id")
        file_url_wix = file_info.get("url")

        # Convert GUID subdomain to public domain
        file_url_wix = file_url_wix.replace(f"https://{WIX_SITE_ID}.usrfiles.com/ugd/", "https://www.dmlsclub.com/_files/ugd/")

        # Create WIX document reference for CMS
        wix_doc_ref = f"wix:document://v1/ugd/{file_id}/{display_name}"

        return {
            "fileId": file_id,
            "wixDocRef": wix_doc_ref,
            "url": file_url_wix,
            "displayName": display_name
        }

    return None


def create_wix_cms_entry(title, wix_doc_ref, publish_date, summary):
    """Create entry in Newsletters CMS collection"""
    endpoint = "https://www.wixapis.com/wix-data/v2/items"

    headers = {
        "Authorization": WIX_API_KEY,
        "wix-site-id": WIX_SITE_ID,
        "Content-Type": "application/json"
    }

    # Format publish date as "Nov 18, 2025"
    if publish_date:
        date_obj = datetime.fromisoformat(publish_date.replace('Z', '+00:00')) if isinstance(publish_date, str) else publish_date
        formatted_date = date_obj.strftime("%b %d, %Y")
    else:
        formatted_date = None

    data = {
        "dataCollectionId": "Newsletters",
        "dataItem": {
            "data": {
                "title": title,
                "newsletter": wix_doc_ref,  # WIX document reference
                "newsletterSummary": summary
            }
        }
    }

    # Only add publishMonth if we have a date
    if formatted_date:
        data["dataItem"]["data"]["publishMonth"] = formatted_date

    response = requests.post(endpoint, headers=headers, json=data)

    if response.status_code == 200:
        result = response.json()
        return result.get("dataItem", {}).get("id")

    return None


# ============================================================
# MAILCHIMP FUNCTIONS
# ============================================================

def read_template():
    """Read the newsletter HTML template."""
    if not TEMPLATE_FILE.exists():
        raise Exception(f"Template file not found: {TEMPLATE_FILE}")

    with open(TEMPLATE_FILE, 'r') as f:
        return f.read()


def create_campaign(month, year):
    """Create a new campaign with title and subject."""
    url = f"{MAILCHIMP_BASE_URL}/campaigns"

    data = {
        "type": "regular",
        "recipients": {
            "list_id": MAILCHIMP_LIST_ID
        },
        "settings": {
            "subject_line": f"DMSC {month} Newsletter is available!",
            "title": f"{month} Newsletter",
            "from_name": "DMSC",
            "reply_to": "dmscnews@gmail.com"
        }
    }

    response = requests.post(url, auth=MAILCHIMP_AUTH, headers=HEADERS, json=data)
    response.raise_for_status()

    return response.json()


def set_campaign_content(campaign_id, html_content):
    """Set the HTML content for a campaign."""
    url = f"{MAILCHIMP_BASE_URL}/campaigns/{campaign_id}/content"

    data = {"html": html_content}

    response = requests.put(url, auth=MAILCHIMP_AUTH, headers=HEADERS, json=data)
    response.raise_for_status()

    return response.json()


def get_campaign_web_url(campaign_id):
    """Get the MailChimp web UI URL for the campaign."""
    url = f"{MAILCHIMP_BASE_URL}/campaigns/{campaign_id}"
    response = requests.get(url, auth=MAILCHIMP_AUTH, headers=HEADERS)
    response.raise_for_status()

    campaign = response.json()
    web_id = campaign.get("web_id")

    # MailChimp campaign editor URL format
    return f"https://{DC}.admin.mailchimp.com/campaigns/edit?id={web_id}"


def main():
    """Main workflow for creating newsletter campaign."""
    print("=" * 60)
    print("DMSC Newsletter Campaign Automation")
    print("Phase 3: Complete Dropbox → WIX → MailChimp Workflow")
    print("=" * 60)
    print()

    # Get month from command line or interactive input
    if len(sys.argv) >= 2:
        month_name = sys.argv[1]
    else:
        month_name = input("Enter newsletter month (e.g., 'December'): ").strip()
        if not month_name:
            print("✗ Month is required")
            sys.exit(1)

    current_year = datetime.now().year
    print(f"Processing: {month_name} {current_year}")
    print()

    # ========================================
    # STEP 1: Find PDF in Dropbox
    # ========================================
    print("Step 1: Finding newsletter PDF in Dropbox...")
    pdf_info = find_newsletter_pdf(month_name)

    if not pdf_info:
        print(f"✗ Could not find newsletter PDF for {month_name}")
        print(f"  Expected location: /Newsletter/Monthly Newsletters/{current_year} Newsletter/{month_name}/")
        sys.exit(1)

    pdf_path = pdf_info["path"]
    pdf_filename = pdf_info["filename"]
    print(f"✓ Found PDF: {pdf_filename}")
    print()

    # ========================================
    # STEP 2: Create Dropbox share link
    # ========================================
    print("Step 2: Creating Dropbox share link...")
    share_link = create_dropbox_share_link(pdf_path)

    if not share_link:
        print("✗ Failed to create Dropbox share link")
        sys.exit(1)

    print(f"✓ Share link created")
    print()

    # ========================================
    # STEP 3: Extract meeting info from Ted's Thoughts
    # ========================================
    print("Step 3: Extracting meeting info from Ted's Thoughts...")
    ted_path = find_teds_thoughts(month_name)

    if ted_path:
        print(f"  Found: {ted_path}")
        docx_bytes = download_dropbox_file(ted_path)
        if docx_bytes:
            meeting_summary = extract_meeting_info(docx_bytes)
            print(f"✓ Meeting summary: {meeting_summary[:60]}...")
        else:
            print("  Warning: Could not download Ted's Thoughts")
            meeting_summary = "See newsletter for meeting details"
    else:
        print("  No Ted's Thoughts found, using default summary")
        meeting_summary = "See newsletter for meeting details"
    print()

    # ========================================
    # STEP 4: Import to WIX Media Manager
    # ========================================
    print("Step 4: Importing PDF to WIX Media Manager...")
    wix_result = import_file_to_wix(share_link, pdf_filename)

    if not wix_result:
        print("✗ Failed to import PDF to WIX")
        sys.exit(1)

    wix_url = wix_result["url"]
    print(f"✓ File imported to WIX")
    print(f"  Public URL: {wix_url}")
    print()

    # ========================================
    # STEP 5: Create WIX CMS entry
    # ========================================
    print("Step 5: Creating WIX CMS entry...")
    cms_title = f"{month_name} {current_year}"

    cms_id = create_wix_cms_entry(
        title=cms_title,
        wix_doc_ref=wix_result["wixDocRef"],
        publish_date=None,  # Leave empty - filled manually in WIX UI
        summary=meeting_summary
    )

    if not cms_id:
        print("✗ Failed to create WIX CMS entry")
        print("  Note: PDF was uploaded successfully to WIX Media Manager")
        sys.exit(1)

    print(f"✓ CMS entry created")
    print(f"  Title: {cms_title}")
    print(f"  Summary: {meeting_summary[:60]}...")
    print()

    # ========================================
    # STEP 6: Read and customize email template
    # ========================================
    print("Step 6: Customizing email template...")
    try:
        template_html = read_template()
        customized_html = template_html.replace("{{MONTH}}", month_name)
        customized_html = customized_html.replace("{{WIX_LINK}}", wix_url)
        print(f"✓ Template customized")
        print()
    except Exception as e:
        print(f"✗ Error with email template: {e}")
        sys.exit(1)

    # ========================================
    # STEP 7: Create MailChimp campaign
    # ========================================
    print("Step 7: Creating MailChimp campaign...")
    try:
        new_campaign = create_campaign(month_name, current_year)
        campaign_id = new_campaign["id"]
        print(f"✓ Campaign created")
        print(f"  Campaign ID: {campaign_id}")
        print(f"  Subject: DMSC {month_name} Newsletter is available!")
        print()
    except Exception as e:
        print(f"✗ Error creating campaign: {e}")
        sys.exit(1)

    # ========================================
    # STEP 8: Upload email content
    # ========================================
    print("Step 8: Uploading email content to MailChimp...")
    try:
        set_campaign_content(campaign_id, customized_html)
        print(f"✓ Email content uploaded")
        print()
    except Exception as e:
        print(f"✗ Error uploading content: {e}")
        sys.exit(1)

    # ========================================
    # SUCCESS SUMMARY
    # ========================================
    print("=" * 60)
    print("SUCCESS! Complete automation workflow finished")
    print("=" * 60)
    print()
    print(f"Newsletter: {month_name} {current_year}")
    print(f"WIX URL: {wix_url}")
    print(f"Meeting Summary: {meeting_summary[:60]}...")
    print()
    print("Next steps:")
    print("1. Review the MailChimp campaign:")
    web_url = get_campaign_web_url(campaign_id)
    print(f"   {web_url}")
    print()
    print("2. Verify:")
    print(f"   - WIX CMS entry at dmlsclub.com/newsletters")
    print(f"   - Email content looks correct")
    print(f"   - WIX link works: {wix_url[:60]}...")
    print()
    print("3. When ready, click 'Send' in MailChimp")
    print()
    print("=" * 60)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nCancelled by user")
        sys.exit(0)
    except Exception as e:
        print(f"\n✗ Unexpected error: {e}")
        sys.exit(1)
