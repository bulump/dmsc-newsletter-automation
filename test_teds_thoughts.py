#!/usr/bin/env python3
"""
Test script to extract meeting info from Ted's Thoughts
"""

import requests
import os
import io
from dotenv import load_dotenv
from docx import Document

load_dotenv()

DROPBOX_ACCESS_TOKEN = os.getenv("DROPBOX_ACCESS_TOKEN")


def find_teds_thoughts(month, year=2025):
    """Find Ted's Thoughts document in Dropbox"""
    folder_path = f"/Newsletter/Monthly Newsletters/{year} Newsletter/{month}"

    url = "https://api.dropboxapi.com/2/files/list_folder"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {"path": folder_path}

    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        entries = response.json().get("entries", [])
        # Find file with "Ted" in the name
        ted_files = [e for e in entries if "ted" in e["name"].lower() and e["name"].endswith(".docx")]

        if ted_files:
            return ted_files[0]["path_lower"]

    return None


def download_dropbox_file(file_path):
    """Download a file from Dropbox"""
    url = "https://content.dropboxapi.com/2/files/download"
    headers = {
        "Authorization": f"Bearer {DROPBOX_ACCESS_TOKEN}",
        "Dropbox-API-Arg": f'{{"path": "{file_path}"}}'
    }

    response = requests.post(url, headers=headers)

    if response.status_code == 200:
        return response.content
    else:
        return None


def extract_meeting_info(docx_bytes):
    """Extract meeting information from Ted's Thoughts document"""
    try:
        doc = Document(io.BytesIO(docx_bytes))

        # Read first few paragraphs looking for meeting information
        meeting_info = []
        for i, para in enumerate(doc.paragraphs[:10]):  # Check first 10 paragraphs
            text = para.text.strip()
            if text and ("meeting" in text.lower() or "speaker" in text.lower() or "program" in text.lower()):
                meeting_info.append(text)

        if meeting_info:
            # Return the first relevant paragraph
            return meeting_info[0]
        else:
            # Fallback: return first non-empty paragraph
            for para in doc.paragraphs[:5]:
                if para.text.strip():
                    return para.text.strip()

        return "See newsletter for meeting details"

    except Exception as e:
        print(f"Error extracting meeting info: {e}")
        return "See newsletter for meeting details"


if __name__ == "__main__":
    import sys

    month = sys.argv[1] if len(sys.argv) > 1 else "November"

    print("=" * 60)
    print(f"Testing Ted's Thoughts Extraction: {month}")
    print("=" * 60)

    # Find Ted's Thoughts
    print(f"\nFinding Ted's Thoughts for {month}...")
    ted_path = find_teds_thoughts(month)

    if not ted_path:
        print(f"✗ Could not find Ted's Thoughts document")
        exit(1)

    print(f"✓ Found: {ted_path}")

    # Download it
    print(f"\nDownloading document...")
    docx_bytes = download_dropbox_file(ted_path)

    if not docx_bytes:
        print("✗ Failed to download")
        exit(1)

    print(f"✓ Downloaded ({len(docx_bytes)} bytes)")

    # Extract meeting info
    print(f"\nExtracting meeting information...")
    meeting_summary = extract_meeting_info(docx_bytes)

    print("\n" + "=" * 60)
    print("Meeting Summary:")
    print("=" * 60)
    print(meeting_summary)
    print("=" * 60)
